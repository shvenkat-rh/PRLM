"""
Document generator for creating Word documents from PR analysis results.
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import os
import logging
from datetime import datetime
from typing import Dict, Any, List

logger = logging.getLogger(__name__)

class DocumentGenerator:
    """Generate Word documents from PR analysis results"""
    
    def __init__(self):
        """Initialize document generator"""
        self.document = None
    
    def create_pr_analysis_report(self, 
                                pr_data, 
                                timeline_metrics, 
                                analysis_result, 
                                file_analysis: Dict = None,
                                conversation_analysis: Dict = None,
                                output_path: str = None) -> str:
        """
        Create a comprehensive PR analysis report in Word format
        
        Args:
            pr_data: PR data from GitHub API
            timeline_metrics: Timeline analysis results
            analysis_result: LLM analysis results
            file_analysis: Detailed file analysis (optional)
            output_path: Output file path (optional)
            
        Returns:
            Path to generated document
        """
        
        # Create new document
        self.document = Document()
        
        # Set up document styles
        self._setup_document_styles()
        
        # Add title page
        self._add_title_page(pr_data)
        
        # Add executive summary
        self._add_executive_summary(pr_data, timeline_metrics, analysis_result)
        
        # Add PR overview
        self._add_pr_overview(pr_data, timeline_metrics)
        
        # Add analysis sections
        self._add_analysis_sections(analysis_result)
        
        # Add detailed file analysis if available
        if file_analysis:
            self._add_file_analysis(file_analysis)
        
        # Add timeline and metrics
        self._add_timeline_analysis(timeline_metrics)
        
        # Add comments analysis
        self._add_comments_analysis(pr_data, conversation_analysis)
        
        # Add appendix
        self._add_appendix(pr_data)
        
        # Generate output path if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_title = "".join(c for c in pr_data.title if c.isalnum() or c in (' ', '-', '_')).rstrip()[:50]
            output_path = f"output/PR_Analysis_{safe_title}_{timestamp}.docx"
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save document
        self.document.save(output_path)
        logger.info(f"PR analysis report saved to {output_path}")
        
        return output_path
    
    def _setup_document_styles(self):
        """Set up custom styles for the document"""
        styles = self.document.styles
        
        # Title style
        if 'Custom Title' not in [style.name for style in styles]:
            title_style = styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(12)
        
        # Heading styles
        if 'Custom Heading 1' not in [style.name for style in styles]:
            h1_style = styles.add_style('Custom Heading 1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Calibri'
            h1_font.size = Pt(18)
            h1_font.bold = True
            h1_style.paragraph_format.space_before = Pt(12)
            h1_style.paragraph_format.space_after = Pt(6)
        
        if 'Custom Heading 2' not in [style.name for style in styles]:
            h2_style = styles.add_style('Custom Heading 2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Calibri'
            h2_font.size = Pt(14)
            h2_font.bold = True
            h2_style.paragraph_format.space_before = Pt(6)
            h2_style.paragraph_format.space_after = Pt(3)
    
    def _add_title_page(self, pr_data):
        """Add title page to document"""
        # Main title
        title = self.document.add_paragraph(f"Pull Request Analysis Report", style='Custom Title')
        
        # Subtitle
        subtitle = self.document.add_paragraph(f"{pr_data.title}", style='Heading 2')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # PR details
        self.document.add_paragraph()  # Space
        details_table = self.document.add_table(rows=6, cols=2)
        details_table.style = 'Table Grid'
        
        details = [
            ("Repository", pr_data.url.split('/pull/')[0].replace('https://github.com/', '')),
            ("PR Number", pr_data.url.split('/pull/')[-1]),
            ("Author", pr_data.author),
            ("State", pr_data.state.upper()),
            ("Created", pr_data.created_at[:10]),
            ("Report Generated", datetime.now().strftime("%Y-%m-%d %H:%M"))
        ]
        
        for i, (label, value) in enumerate(details):
            details_table.rows[i].cells[0].text = label
            details_table.rows[i].cells[1].text = str(value)
            details_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Page break
        self.document.add_page_break()
    
    def _add_executive_summary(self, pr_data, timeline_metrics, analysis_result):
        """Add executive summary section"""
        self.document.add_heading('Executive Summary', level=1)
        
        # Key metrics summary
        summary_para = self.document.add_paragraph()
        summary_para.add_run("This report provides a comprehensive analysis of Pull Request: ").bold = True
        summary_para.add_run(f'"{pr_data.title}" ')
        summary_para.add_run("submitted by ").bold = True
        summary_para.add_run(f"{pr_data.author}. ")
        
        # Add timeline summary
        if timeline_metrics.time_to_first_review:
            summary_para.add_run(f"The PR received its first review after {timeline_metrics.time_to_first_review:.1f} hours. ")
        
        if timeline_metrics.time_to_merge:
            summary_para.add_run(f"It was merged after {timeline_metrics.time_to_merge:.1f} hours total. ")
        
        summary_para.add_run(f"The PR involved {len(pr_data.changed_files)} file(s) with {len(pr_data.comments)} comments and {timeline_metrics.review_cycles} review cycle(s).")
        
        # Add LLM summary if available
        if analysis_result and analysis_result.summary:
            self.document.add_paragraph()
            summary_heading = self.document.add_paragraph()
            summary_heading.add_run("AI Analysis Summary:").bold = True
            self.document.add_paragraph(analysis_result.summary)
    
    def _add_pr_overview(self, pr_data, timeline_metrics):
        """Add PR overview section"""
        self.document.add_heading('Pull Request Overview', level=1)
        
        # Basic information
        self.document.add_heading('Basic Information', level=2)
        
        info_table = self.document.add_table(rows=8, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("Title", pr_data.title),
            ("Author", pr_data.author),
            ("State", pr_data.state),
            ("Created", pr_data.created_at),
            ("Last Updated", pr_data.updated_at),
            ("Files Changed", len(pr_data.changed_files)),
            ("Total Comments", len(pr_data.comments)),
            ("Review Cycles", timeline_metrics.review_cycles)
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
            info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Description
        if pr_data.description:
            self.document.add_heading('Description', level=2)
            self.document.add_paragraph(pr_data.description[:1000] + ("..." if len(pr_data.description) > 1000 else ""))
    
    def _add_analysis_sections(self, analysis_result):
        """Add LLM analysis sections"""
        if not analysis_result:
            return
        
        # Key Changes
        if analysis_result.key_changes:
            self.document.add_heading('Key Changes', level=1)
            for change in analysis_result.key_changes:
                para = self.document.add_paragraph(style='List Bullet')
                para.text = change
        
        # Developer Mistakes
        if analysis_result.developer_mistakes:
            self.document.add_heading('Developer Mistakes & Learning Opportunities', level=1)
            for mistake in analysis_result.developer_mistakes:
                para = self.document.add_paragraph(style='List Bullet')
                para.text = mistake
        
        # Code Quality Issues
        if analysis_result.code_quality_issues:
            self.document.add_heading('Code Quality Issues', level=1)
            for issue in analysis_result.code_quality_issues:
                para = self.document.add_paragraph(style='List Bullet')
                para.text = issue
        
        # Suggestions
        if analysis_result.suggestions:
            self.document.add_heading('Suggestions & Recommendations', level=1)
            for suggestion in analysis_result.suggestions:
                para = self.document.add_paragraph(style='List Bullet')
                para.text = suggestion
        
        # Comment Analysis
        if analysis_result.comment_analysis:
            self.document.add_heading('Comment & Review Analysis', level=1)
            self.document.add_paragraph(analysis_result.comment_analysis)
        
        # Overall Assessment
        if analysis_result.overall_assessment:
            self.document.add_heading('Overall Assessment', level=1)
            self.document.add_paragraph(analysis_result.overall_assessment)
    
    def _add_file_analysis(self, file_analysis):
        """Add detailed file analysis section"""
        if not file_analysis:
            return
            
        self.document.add_heading('Detailed File Analysis', level=1)
        
        for filename, analysis in file_analysis.items():
            self.document.add_heading(filename, level=2)
            
            # Add file metrics
            metrics_para = self.document.add_paragraph()
            metrics_para.add_run("Complexity Score: ").bold = True
            metrics_para.add_run(f"{analysis.get('complexity_score', 'N/A'):.1f}/10")
            metrics_para.add_run(" | Risk Level: ").bold = True
            risk_level = analysis.get('risk_level', 'unknown').upper()
            risk_run = metrics_para.add_run(risk_level)
            
            # Color code risk levels
            if risk_level == 'HIGH':
                risk_run.font.color.rgb = None  # Red would require more complex color handling
            
            # Add analysis text
            if analysis.get('analysis'):
                self.document.add_paragraph(analysis['analysis'])
    
    def _add_timeline_analysis(self, timeline_metrics):
        """Add timeline analysis section"""
        self.document.add_heading('Timeline Analysis', level=1)
        
        # Timeline table
        timeline_table = self.document.add_table(rows=6, cols=2)
        timeline_table.style = 'Table Grid'
        
        timeline_data = [
            ("Time to First Review", f"{timeline_metrics.time_to_first_review:.1f} hours" if timeline_metrics.time_to_first_review else "No reviews"),
            ("Time to Merge", f"{timeline_metrics.time_to_merge:.1f} hours" if timeline_metrics.time_to_merge else "Not merged"),
            ("Total Lifecycle", f"{timeline_metrics.total_lifecycle:.1f} hours"),
            ("Review Cycles", str(timeline_metrics.review_cycles)),
            ("Comment Frequency", f"{timeline_metrics.comment_frequency:.1f} comments/day"),
            ("First Review Date", timeline_metrics.first_review_at.strftime("%Y-%m-%d %H:%M") if timeline_metrics.first_review_at else "N/A")
        ]
        
        for i, (label, value) in enumerate(timeline_data):
            timeline_table.rows[i].cells[0].text = label
            timeline_table.rows[i].cells[1].text = value
            timeline_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    def _add_comments_analysis(self, pr_data, conversation_analysis=None):
        """Add comprehensive comments and conversation analysis section"""
        if not pr_data.comments:
            return
            
        self.document.add_heading('Comments & Discussions Analysis', level=1)
        
        # Comments summary
        comment_types = {}
        for comment in pr_data.comments:
            comment_types[comment.type] = comment_types.get(comment.type, 0) + 1
        
        summary_para = self.document.add_paragraph("Comment Summary: ")
        for comment_type, count in comment_types.items():
            summary_para.add_run(f"{comment_type.replace('_', ' ').title()}: {count}; ")
        
        # Add conversation analysis if available
        if conversation_analysis:
            self._add_conversation_threads_analysis(conversation_analysis)
            self._add_reviewer_profiles_analysis(conversation_analysis)
        
        # Chronological conversation flow
        self.document.add_heading('Conversation Timeline', level=2)
        sorted_comments = sorted(pr_data.comments, key=lambda x: x.created_at)
        
        for i, comment in enumerate(sorted_comments):
            comment_para = self.document.add_paragraph()
            
            # Add timestamp and author
            timestamp = comment.created_at[:16].replace('T', ' ')
            comment_para.add_run(f"{i+1}. [{timestamp}] {comment.author} ({comment.type}): ").bold = True
            
            # Add comment body
            comment_text = comment.body.strip()
            if len(comment_text) > 400:
                comment_text = comment_text[:400] + "..."
            comment_para.add_run(comment_text)
            
            # Add spacing between comments
            self.document.add_paragraph()
    
    def _add_conversation_threads_analysis(self, conversation_analysis):
        """Add detailed conversation threads analysis"""
        if not conversation_analysis.get('threads'):
            return
            
        self.document.add_heading('Conversation Threads Analysis', level=2)
        
        threads = conversation_analysis['threads']
        metrics = conversation_analysis['metrics']
        
        # Thread summary
        summary_para = self.document.add_paragraph()
        summary_para.add_run(f"Total Conversation Threads: {metrics.total_threads}").bold = True
        summary_para.add_run(f"\nResolved Threads: {metrics.resolved_threads}")
        summary_para.add_run(f"\nUnresolved Threads: {metrics.unresolved_threads}")
        summary_para.add_run(f"\nAverage Responses per Thread: {metrics.avg_responses_per_thread:.1f}")
        summary_para.add_run(f"\nCommunication Tone: {metrics.communication_tone.title()}")
        
        # Individual thread analysis
        self.document.add_heading('Discussion Threads', level=3)
        
        # Sort threads by number of comments (most active first)
        sorted_threads = sorted(threads, key=lambda t: len(t.comments), reverse=True)
        
        for i, thread in enumerate(sorted_threads[:10], 1):  # Top 10 threads
            thread_heading = self.document.add_paragraph()
            thread_heading.add_run(f"Thread {i}: {thread.topic}").bold = True
            
            thread_info = self.document.add_paragraph()
            thread_info.add_run(f"Type: {thread.thread_type.replace('_', ' ').title()}")
            thread_info.add_run(f" | Participants: {', '.join(thread.participants)}")
            thread_info.add_run(f" | Comments: {len(thread.comments)}")
            thread_info.add_run(f" | Status: {thread.resolution_status.title()}")
            
            # Show conversation flow for this thread
            if len(thread.comments) <= 5:  # Show all comments for short threads
                for comment in thread.comments:
                    conv_para = self.document.add_paragraph(style='List Bullet')
                    conv_para.add_run(f"{comment.author}: ").bold = True
                    conv_para.add_run(comment.body[:200] + ("..." if len(comment.body) > 200 else ""))
            else:  # Show first and last comments for long threads
                conv_para = self.document.add_paragraph(style='List Bullet')
                first_comment = thread.comments[0]
                conv_para.add_run(f"{first_comment.author} (initial): ").bold = True
                conv_para.add_run(first_comment.body[:200] + ("..." if len(first_comment.body) > 200 else ""))
                
                if len(thread.comments) > 2:
                    conv_para = self.document.add_paragraph()
                    conv_para.add_run(f"... {len(thread.comments) - 2} more exchanges ...").italic = True
                
                conv_para = self.document.add_paragraph(style='List Bullet')
                last_comment = thread.comments[-1]
                conv_para.add_run(f"{last_comment.author} (final): ").bold = True
                conv_para.add_run(last_comment.body[:200] + ("..." if len(last_comment.body) > 200 else ""))
            
            # Add spacing between threads
            self.document.add_paragraph()
    
    def _add_reviewer_profiles_analysis(self, conversation_analysis):
        """Add reviewer engagement and behavior analysis"""
        if not conversation_analysis.get('reviewer_profiles'):
            return
            
        self.document.add_heading('Reviewer Engagement Analysis', level=2)
        
        profiles = conversation_analysis['reviewer_profiles']
        
        # Reviewer summary table
        if profiles:
            reviewer_table = self.document.add_table(rows=len(profiles) + 1, cols=6)
            reviewer_table.style = 'Table Grid'
            
            # Header row
            header_cells = reviewer_table.rows[0].cells
            headers = ["Reviewer", "Total Comments", "Questions", "Suggestions", "Approvals", "Engagement"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for i, (name, profile) in enumerate(profiles.items(), 1):
                row_cells = reviewer_table.rows[i].cells
                row_cells[0].text = name
                row_cells[1].text = str(profile.total_comments)
                row_cells[2].text = str(profile.comment_types.get('questions', 0))
                row_cells[3].text = str(profile.comment_types.get('suggestions', 0))
                row_cells[4].text = str(profile.comment_types.get('approvals', 0))
                row_cells[5].text = profile.engagement_level.title()
        
        # Communication patterns analysis
        patterns = conversation_analysis.get('patterns', {})
        if patterns:
            self.document.add_heading('Communication Patterns', level=3)
            
            patterns_para = self.document.add_paragraph()
            patterns_para.add_run("Author Responsiveness: ").bold = True
            responsiveness = patterns.get('author_responsiveness', 0)
            patterns_para.add_run(f"{responsiveness:.1f} responses per reviewer comment")
            
            if patterns.get('collaboration_indicators'):
                patterns_para.add_run("\nCollaborative Indicators: ").bold = True
                collab_count = len(set(patterns['collaboration_indicators']))
                patterns_para.add_run(f"{collab_count} instances of positive collaboration")
            
            if patterns.get('conflict_indicators'):
                patterns_para.add_run("\nDiscussion Points: ").bold = True
                conflict_count = len(set(patterns['conflict_indicators']))
                patterns_para.add_run(f"{conflict_count} instances of differing viewpoints")
        
        # Conversation summary
        summary = conversation_analysis.get('summary', '')
        if summary:
            self.document.add_heading('Conversation Summary', level=3)
            self.document.add_paragraph(summary)
    
    def _add_appendix(self, pr_data):
        """Add appendix with detailed data"""
        self.document.add_page_break()
        self.document.add_heading('Appendix', level=1)
        
        # File changes summary
        if pr_data.changed_files:
            self.document.add_heading('Changed Files Summary', level=2)
            
            files_table = self.document.add_table(rows=len(pr_data.changed_files) + 1, cols=5)
            files_table.style = 'Table Grid'
            
            # Header
            header_cells = files_table.rows[0].cells
            headers = ["Filename", "Status", "Additions", "Deletions", "Total Changes"]
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for i, file_info in enumerate(pr_data.changed_files, 1):
                row_cells = files_table.rows[i].cells
                row_cells[0].text = file_info['filename']
                row_cells[1].text = file_info['status']
                row_cells[2].text = str(file_info['additions'])
                row_cells[3].text = str(file_info['deletions'])
                row_cells[4].text = str(file_info['changes'])
        
        # Commits summary
        if pr_data.commits:
            self.document.add_heading('Commits Summary', level=2)
            
            for commit in pr_data.commits[-10:]:  # Last 10 commits
                commit_para = self.document.add_paragraph()
                commit_para.add_run(f"{commit['sha'][:8]} ").bold = True
                commit_para.add_run(f"by {commit['author']} ({commit['date'][:10]}): ")
                commit_para.add_run(commit['message'][:100] + ("..." if len(commit['message']) > 100 else ""))
    
    def create_combined_pr_analysis_report(self, 
                                         all_pr_data: List, 
                                         all_timeline_metrics: List, 
                                         all_analysis_results: List, 
                                         all_file_analyses: List,
                                         all_conversation_analyses: List,
                                         output_path: str = None) -> str:
        """
        Create a combined analysis report for multiple PRs
        
        Args:
            all_pr_data: List of PR data objects
            all_timeline_metrics: List of timeline metrics
            all_analysis_results: List of LLM analysis results
            all_file_analyses: List of detailed file analyses
            all_conversation_analyses: List of conversation analyses
            output_path: Custom output path for the Word document
            
        Returns:
            Path to generated combined document
        """
        
        # Create new document
        self.document = Document()
        
        # Set up document styles
        self._setup_document_styles()
        
        # Add combined title page
        self._add_combined_title_page(all_pr_data)
        
        # Add executive summary for all PRs
        self._add_combined_executive_summary(all_pr_data, all_timeline_metrics, all_analysis_results)
        
        # Add comparative analysis
        self._add_comparative_analysis(all_pr_data, all_timeline_metrics, all_analysis_results)
        
        # Add individual PR sections
        for i, (pr_data, timeline_metrics, analysis_result, file_analysis, conversation_analysis) in enumerate(
            zip(all_pr_data, all_timeline_metrics, all_analysis_results, all_file_analyses, all_conversation_analyses), 1
        ):
            # Add page break before each PR (except first)
            if i > 1:
                self.document.add_page_break()
            
            # Add PR-specific analysis
            self._add_individual_pr_section(pr_data, timeline_metrics, analysis_result, file_analysis, conversation_analysis, i)
        
        # Add combined insights and recommendations
        self._add_combined_insights(all_pr_data, all_timeline_metrics, all_analysis_results, all_conversation_analyses)
        
        # Generate output path if not provided
        if not output_path:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = f"analysis_outputs/pr_summaries/Combined_PR_Analysis_{len(all_pr_data)}_PRs_{timestamp}.docx"
        
        # Ensure output directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        # Save document
        self.document.save(output_path)
        logger.info(f"Combined PR analysis report saved to {output_path}")
        
        return output_path
    
    def _add_combined_title_page(self, all_pr_data):
        """Add title page for combined analysis"""
        # Main title
        title = self.document.add_paragraph(f"Combined Pull Request Analysis Report", style='Custom Title')
        
        # Subtitle
        subtitle = self.document.add_paragraph(f"Analysis of {len(all_pr_data)} Pull Requests", style='Heading 2')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # PR overview table
        self.document.add_paragraph()  # Space
        overview_table = self.document.add_table(rows=len(all_pr_data) + 1, cols=4)
        overview_table.style = 'Table Grid'
        
        # Header
        header_cells = overview_table.rows[0].cells
        headers = ["PR", "Title", "Author", "Repository"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # PR details
        for i, pr_data in enumerate(all_pr_data, 1):
            row_cells = overview_table.rows[i].cells
            row_cells[0].text = f"PR {i}"
            row_cells[1].text = pr_data.title[:50] + ("..." if len(pr_data.title) > 50 else "")
            row_cells[2].text = pr_data.author
            row_cells[3].text = pr_data.url.split('/pull/')[0].replace('https://github.com/', '')
        
        # Generation details
        self.document.add_paragraph()
        details_table = self.document.add_table(rows=2, cols=2)
        details_table.style = 'Table Grid'
        
        details = [
            ("Total PRs Analyzed", str(len(all_pr_data))),
            ("Report Generated", datetime.now().strftime("%Y-%m-%d %H:%M"))
        ]
        
        for i, (label, value) in enumerate(details):
            details_table.rows[i].cells[0].text = label
            details_table.rows[i].cells[1].text = value
            details_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Page break
        self.document.add_page_break()
    
    def _add_combined_executive_summary(self, all_pr_data, all_timeline_metrics, all_analysis_results):
        """Add executive summary for all PRs combined"""
        self.document.add_heading('Executive Summary', level=1)
        
        # Combined metrics
        total_files = sum(len(pr.changed_files) for pr in all_pr_data)
        total_comments = sum(len(pr.comments) for pr in all_pr_data)
        total_commits = sum(len(pr.commits) for pr in all_pr_data)
        
        # Average timeline metrics
        avg_time_to_review = sum(tm.time_to_first_review or 0 for tm in all_timeline_metrics) / len(all_timeline_metrics)
        merged_prs = [tm for tm in all_timeline_metrics if tm.time_to_merge]
        avg_time_to_merge = sum(tm.time_to_merge for tm in merged_prs) / max(1, len(merged_prs))
        
        summary_para = self.document.add_paragraph()
        summary_para.add_run(f"This report analyzes {len(all_pr_data)} Pull Requests ").bold = True
        summary_para.add_run(f"involving {total_files} file changes, {total_comments} comments, and {total_commits} commits. ")
        
        if avg_time_to_review > 0:
            summary_para.add_run(f"The average time to first review was {avg_time_to_review:.1f} hours. ")
        
        if len(merged_prs) > 0:
            summary_para.add_run(f"Among the {len(merged_prs)} merged PRs, the average time to merge was {avg_time_to_merge:.1f} hours.")
        
        # Combined insights from AI analysis
        if any(all_analysis_results):
            self.document.add_paragraph()
            insights_heading = self.document.add_paragraph()
            insights_heading.add_run("Key Patterns Across All PRs:").bold = True
            
            # Aggregate common themes
            all_mistakes = []
            all_suggestions = []
            for result in all_analysis_results:
                if result:
                    all_mistakes.extend(result.developer_mistakes)
                    all_suggestions.extend(result.suggestions)
            
            if all_mistakes:
                self.document.add_paragraph()
                mistakes_heading = self.document.add_paragraph()
                mistakes_heading.add_run("Common Learning Opportunities:").bold = True
                for mistake in all_mistakes[:5]:  # Top 5
                    mistake_para = self.document.add_paragraph(style='List Bullet')
                    mistake_para.text = mistake
    
    def _add_comparative_analysis(self, all_pr_data, all_timeline_metrics, all_analysis_results):
        """Add comparative analysis section"""
        self.document.add_heading('Comparative Analysis', level=1)
        
        # Timeline comparison table
        self.document.add_heading('Timeline Comparison', level=2)
        
        timeline_table = self.document.add_table(rows=len(all_pr_data) + 1, cols=7)
        timeline_table.style = 'Table Grid'
        
        # Header
        header_cells = timeline_table.rows[0].cells
        headers = ["PR", "Author", "Files", "Comments", "Time to Review (hrs)", "Time to Merge (hrs)", "Review Cycles"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Data rows
        for i, (pr_data, timeline_metrics) in enumerate(zip(all_pr_data, all_timeline_metrics), 1):
            row_cells = timeline_table.rows[i].cells
            row_cells[0].text = f"PR {i}"
            row_cells[1].text = pr_data.author
            row_cells[2].text = str(len(pr_data.changed_files))
            row_cells[3].text = str(len(pr_data.comments))
            row_cells[4].text = f"{timeline_metrics.time_to_first_review:.1f}" if timeline_metrics.time_to_first_review else "N/A"
            row_cells[5].text = f"{timeline_metrics.time_to_merge:.1f}" if timeline_metrics.time_to_merge else "N/A"
            row_cells[6].text = str(timeline_metrics.review_cycles)
        
        # Patterns and trends
        self.document.add_heading('Patterns and Trends', level=2)
        
        # Analyze patterns
        patterns_para = self.document.add_paragraph()
        patterns_para.add_run("Analysis of patterns across all PRs:").bold = True
        
        # Calculate some comparative metrics
        avg_files = sum(len(pr.changed_files) for pr in all_pr_data) / len(all_pr_data)
        avg_comments = sum(len(pr.comments) for pr in all_pr_data) / len(all_pr_data)
        
        patterns_para.add_run(f"\n• Average files changed per PR: {avg_files:.1f}")
        patterns_para.add_run(f"\n• Average comments per PR: {avg_comments:.1f}")
        
        # Find patterns by author
        authors = {}
        for pr_data, timeline_metrics in zip(all_pr_data, all_timeline_metrics):
            if pr_data.author not in authors:
                authors[pr_data.author] = []
            authors[pr_data.author].append(timeline_metrics)
        
        if len(authors) > 1:
            patterns_para.add_run(f"\n• Analysis covers {len(authors)} different authors")
    
    def _add_individual_pr_section(self, pr_data, timeline_metrics, analysis_result, file_analysis, conversation_analysis, pr_number):
        """Add individual PR analysis section"""
        
        # PR section header
        self.document.add_heading(f'Pull Request {pr_number}: {pr_data.title}', level=1)
        
        # PR overview
        self.document.add_heading('Overview', level=2)
        
        overview_table = self.document.add_table(rows=8, cols=2)
        overview_table.style = 'Table Grid'
        
        overview_data = [
            ("URL", pr_data.url),
            ("Author", pr_data.author),
            ("State", pr_data.state),
            ("Created", pr_data.created_at[:10]),
            ("Files Changed", len(pr_data.changed_files)),
            ("Comments", len(pr_data.comments)),
            ("Review Cycles", timeline_metrics.review_cycles),
            ("Time to Merge", f"{timeline_metrics.time_to_merge:.1f} hours" if timeline_metrics.time_to_merge else "Not merged")
        ]
        
        for i, (label, value) in enumerate(overview_data):
            overview_table.rows[i].cells[0].text = label
            overview_table.rows[i].cells[1].text = str(value)
            overview_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        
        # Add description if available
        if pr_data.description:
            self.document.add_heading('Description', level=3)
            self.document.add_paragraph(pr_data.description[:500] + ("..." if len(pr_data.description) > 500 else ""))
        
        # Add AI analysis sections for this PR
        if analysis_result:
            if analysis_result.key_changes:
                self.document.add_heading('Key Changes', level=3)
                for change in analysis_result.key_changes[:5]:
                    change_para = self.document.add_paragraph(style='List Bullet')
                    change_para.text = change
        
        # Add conversation analysis for this PR
        if conversation_analysis:
            self._add_conversation_threads_analysis(conversation_analysis)
    
    def _add_combined_insights(self, all_pr_data, all_timeline_metrics, all_analysis_results, all_conversation_analyses):
        """Add combined insights and recommendations section"""
        self.document.add_page_break()
        self.document.add_heading('Combined Insights and Recommendations', level=1)
        
        # Overall team patterns
        self.document.add_heading('Team Development Patterns', level=2)
        
        # Aggregate all conversation analyses
        total_threads = sum(len(conv.get('threads', [])) for conv in all_conversation_analyses if conv)
        total_resolved = sum(sum(1 for t in conv.get('threads', []) if t.resolution_status == 'resolved') 
                           for conv in all_conversation_analyses if conv)
        
        if total_threads > 0:
            resolution_rate = (total_resolved / total_threads) * 100
            patterns_para = self.document.add_paragraph()
            patterns_para.add_run("Team Communication Analysis:").bold = True
            patterns_para.add_run(f"\n• Total conversation threads across all PRs: {total_threads}")
            patterns_para.add_run(f"\n• Overall resolution rate: {resolution_rate:.1f}%")
        
        # Common improvement areas
        self.document.add_heading('Common Improvement Areas', level=2)
        
        all_mistakes = []
        all_suggestions = []
        
        for result in all_analysis_results:
            if result:
                all_mistakes.extend(result.developer_mistakes)
                all_suggestions.extend(result.suggestions)
        
        if all_mistakes:
            mistake_counts = {}
            for mistake in all_mistakes:
                # Simple keyword matching to group similar mistakes
                key_words = mistake.lower().split()[:3]  # First 3 words as key
                key = ' '.join(key_words)
                mistake_counts[key] = mistake_counts.get(key, 0) + 1
            
            # Show most common patterns
            common_mistakes = sorted(mistake_counts.items(), key=lambda x: x[1], reverse=True)[:5]
            
            for pattern, count in common_mistakes:
                if count > 1:  # Only show patterns that appear multiple times
                    mistake_para = self.document.add_paragraph(style='List Bullet')
                    mistake_para.text = f"Pattern: {pattern} (appeared in {count} PRs)"
        
        # Recommendations
        self.document.add_heading('Team Recommendations', level=2)
        
        recommendations = [
            f"Consider implementing automated checks to reduce common review cycles",
            f"Average review time of {sum(tm.time_to_first_review or 0 for tm in all_timeline_metrics) / len(all_timeline_metrics):.1f} hours could be optimized",
            f"Documentation and process improvements based on recurring patterns"
        ]
        
        for rec in recommendations:
            rec_para = self.document.add_paragraph(style='List Bullet')
            rec_para.text = rec 